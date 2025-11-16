#!/usr/bin/env python3
"""
Script for printing daily schedule from CalDAV calendar to Word document.
"""

import os
import platform
import argparse
from datetime import datetime, timedelta, date, timezone
from caldav import DAVClient
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
import vobject
import requests
from requests.auth import HTTPBasicAuth
from urllib.parse import urljoin


def parse_timezone_offset(offset_str):
    """Parse timezone offset string to timezone object.
    
    Args:
        offset_str: String like '+3', '-5', '0', '+5.5' representing hours offset from UTC
    
    Returns:
        timezone object with the specified offset
    """
    try:
        # Default to +3 (Moscow time) if not specified
        if not offset_str:
            offset_str = '+3'
        
        # Parse the offset as float to support half-hour offsets
        offset_hours = float(offset_str)
        
        # Create timezone with the offset
        tz = timezone(timedelta(hours=offset_hours))
        return tz
    except ValueError:
        raise ValueError(
            f"Invalid timezone offset: '{offset_str}'. "
            "Use numeric offset in hours, e.g., +3, -5, 0, +5.5"
        )


def load_config():
    """Load configuration from .env file."""
    load_dotenv()
    
    caldav_username = os.getenv('CALDAV_USERNAME')
    caldav_password = os.getenv('CALDAV_PASSWORD')
    
    # Parse timezone offset
    timezone_offset = os.getenv('TIMEZONE_OFFSET', '+3')
    tz = parse_timezone_offset(timezone_offset)
    
    config = {
        'caldav_url': os.getenv('CALDAV_URL'),
        'caldav_username': caldav_username,
        'caldav_password': caldav_password,
        'carddav_url': os.getenv('CARDDAV_URL'),
        'carddav_username': os.getenv('CARDDAV_USERNAME', caldav_username),
        'carddav_password': os.getenv('CARDDAV_PASSWORD', caldav_password),
        'timezone': tz,
        'document_title': os.getenv('DOCUMENT_TITLE', 'Расписание'),
        'output_path': os.getenv('OUTPUT_PATH', '.'),
        'filename_prefix': os.getenv('FILENAME_PREFIX', 'schedule_')
    }
    
    # Check if all required parameters are present
    if not all([config['caldav_url'], config['caldav_username'], config['caldav_password']]):
        raise ValueError(
            "Missing required configuration. Please check .env file.\n"
            "Required: CALDAV_URL, CALDAV_USERNAME, CALDAV_PASSWORD"
        )
    
    # Ensure output path exists
    if config['output_path'] and config['output_path'] != '.':
        os.makedirs(config['output_path'], exist_ok=True)
    
    return config


def connect_to_calendar(caldav_url, username, password):
    """Connect to CalDAV server and return calendar."""
    try:
        client = DAVClient(
            url=caldav_url,
            username=username,
            password=password
        )
        
        principal = client.principal()
        calendars = principal.calendars()
        
        if not calendars:
            raise ValueError("No calendars found for this user")
        
        # Return the first calendar (you can modify this to select a specific calendar)
        return calendars[0]
    
    except Exception as e:
        raise ConnectionError(f"Failed to connect to CalDAV server: {str(e)}")


def discover_addressbooks(addressbook_home_url, username, password):
    """Discover all available addressbooks from CardDAV addressbook home URL."""
    from xml.etree import ElementTree as ET
    
    try:
        # Direct PROPFIND to addressbook home to get list of addressbooks
        propfind_books = '''<?xml version='1.0'?>
<D:propfind xmlns:D="DAV:">
  <D:prop>
    <D:resourcetype />
    <D:displayname />
  </D:prop>
</D:propfind>'''
        
        response = requests.request(
            'PROPFIND',
            addressbook_home_url,
            headers={'Depth': '1', 'Content-Type': 'application/xml; charset=utf-8'},
            data=propfind_books,
            auth=HTTPBasicAuth(username, password),
            timeout=30
        )
        
        if response.status_code not in [200, 207]:
            print(f"Warning: PROPFIND to addressbook home failed with status {response.status_code}")
            return []
        
        root = ET.fromstring(response.content)
        ns_full = {'D': 'DAV:', 'E': 'urn:ietf:params:xml:ns:carddav'}
        
        addressbooks = []
        for resp in root.findall('D:response', ns_full):
            # Check if it's an addressbook
            resourcetype = resp.find('.//D:resourcetype', ns_full)
            if resourcetype is not None:
                # Look for addressbook resourcetype
                is_addressbook = resourcetype.find('E:addressbook', ns_full) is not None
                
                if is_addressbook:
                    href_elem = resp.find('D:href', ns_full)
                    displayname_elem = resp.find('.//D:displayname', ns_full)
                    
                    if href_elem is not None and href_elem.text:
                        # Build full URL from href
                        href_text = href_elem.text
                        if href_text.startswith('http'):
                            book_url = href_text
                        else:
                            book_url = urljoin(addressbook_home_url, href_text)
                        
                        book_name = displayname_elem.text if displayname_elem is not None and displayname_elem.text else 'Unknown'
                        addressbooks.append({'url': book_url, 'name': book_name})
        
        return addressbooks
    
    except Exception as e:
        print(f"Warning: Failed to discover addressbooks: {str(e)}")
        return []


def load_contacts_from_addressbook(addressbook_url, username, password):
    """Load contacts from a specific addressbook using addressbook-multiget."""
    from xml.etree import ElementTree as ET
    
    email_to_name = {}
    
    try:
        # Step 1: PROPFIND to get list of all contacts (hrefs)
        propfind_body = '''<?xml version='1.0'?>
<D:propfind xmlns:D="DAV:">
  <D:prop>
    <D:getetag/>
    <D:getcontenttype/>
  </D:prop>
</D:propfind>'''
        
        response = requests.request(
            'PROPFIND',
            addressbook_url,
            headers={'Depth': '1', 'Content-Type': 'application/xml; charset=utf-8'},
            data=propfind_body,
            auth=HTTPBasicAuth(username, password),
            timeout=30
        )
        
        if response.status_code not in [200, 207]:
            return email_to_name
        
        root = ET.fromstring(response.content)
        ns = {'D': 'DAV:'}
        
        # Collect all vCard hrefs
        vcard_hrefs = []
        for resp in root.findall('D:response', ns):
            contenttype = resp.find('.//D:getcontenttype', ns)
            if contenttype is not None and contenttype.text and 'vcard' in contenttype.text.lower():
                href_elem = resp.find('D:href', ns)
                if href_elem is not None and href_elem.text:
                    vcard_hrefs.append(href_elem.text)
        
        if not vcard_hrefs:
            return email_to_name
        
        # Step 2: Use addressbook-multiget to fetch all vCards at once
        href_elements = '\n'.join([f'<D:href>{href}</D:href>' for href in vcard_hrefs])
        
        multiget_body = f'''<?xml version="1.0" encoding="utf-8" ?>
   <A:addressbook-multiget xmlns:D="DAV:" xmlns:A="urn:ietf:params:xml:ns:carddav">
     <D:prop>
       <D:getetag/>
       <D:getcontenttype/>
       <A:address-data/>
     </D:prop>
     {href_elements}
   </A:addressbook-multiget>'''
        
        response = requests.request(
            'REPORT',
            addressbook_url,
            headers={'Depth': '1', 'Content-Type': 'application/xml; charset=utf-8'},
            data=multiget_body,
            auth=HTTPBasicAuth(username, password),
            timeout=60
        )
        
        if response.status_code not in [200, 207]:
            print(f"Warning: REPORT addressbook-multiget failed with status {response.status_code}")
            return email_to_name
        
        # Parse response and extract vCards
        root = ET.fromstring(response.content)
        ns_full = {'D': 'DAV:', 'E': 'urn:ietf:params:xml:ns:carddav'}
        
        for resp in root.findall('D:response', ns_full):
            address_data = resp.find('.//E:address-data', ns_full)
            if address_data is not None and address_data.text:
                try:
                    # Parse vCard
                    vcard = vobject.readOne(address_data.text)
                    
                    # Get full name - prioritize structured name (N) over full name (FN)
                    full_name = None
                    
                    # First, try to build name from structured components (Family, Given, Middle/Additional)
                    if hasattr(vcard, 'n'):
                        n = vcard.n.value
                        # Check if at least family or given name is present
                        if n.family or n.given:
                            name_parts = []
                            if n.family:
                                name_parts.append(n.family)
                            if n.given:
                                name_parts.append(n.given)
                            if n.additional:
                                name_parts.append(n.additional)
                            full_name = ' '.join(name_parts).strip()
                    
                    # If structured name is empty, fall back to full name (FN)
                    if not full_name and hasattr(vcard, 'fn') and vcard.fn.value:
                        full_name = str(vcard.fn.value).strip()
                    
                    # Skip contacts without names
                    if not full_name:
                        continue
                    
                    # Get email addresses
                    if hasattr(vcard, 'email_list'):
                        for email in vcard.email_list:
                            email_value = str(email.value).lower().strip()
                            if email_value:
                                email_to_name[email_value] = full_name
                    elif hasattr(vcard, 'email'):
                        email_value = str(vcard.email.value).lower().strip()
                        if email_value:
                            email_to_name[email_value] = full_name
                
                except Exception:
                    # Skip contacts that can't be parsed
                    continue
    
    except Exception as e:
        print(f"Warning: Failed to load contacts from addressbook: {str(e)}")
    
    return email_to_name


def load_contacts_from_carddav(carddav_url, username, password):
    """Load all contacts from all CardDAV addressbooks and create email to name mapping."""
    email_to_name = {}
    
    try:
        # Discover all addressbooks
        addressbooks = discover_addressbooks(carddav_url, username, password)
        
        if not addressbooks:
            print("Warning: No addressbooks found")
            return email_to_name
        
        print(f"Found {len(addressbooks)} addressbook(s): {', '.join([ab['name'] for ab in addressbooks])}")
        
        # Load contacts from each addressbook
        for addressbook in addressbooks:
            print(f"Loading contacts from '{addressbook['name']}'...")
            book_contacts = load_contacts_from_addressbook(
                addressbook['url'],
                username,
                password
            )
            
            # Merge contacts (later addressbooks override earlier ones)
            email_to_name.update(book_contacts)
            print(f"  Loaded {len(book_contacts)} email mappings")
        
        print(f"Total: {len(email_to_name)} unique email mappings loaded")
    
    except Exception as e:
        print(f"Warning: Failed to load contacts from CardDAV: {str(e)}")
    
    return email_to_name


def print_document(filepath):
    """Print document to default printer (Windows only).
    
    Args:
        filepath: Full path to the document to print
    
    Returns:
        bool: True if printing was initiated successfully, False otherwise
    """
    # Check if running on Windows
    if platform.system() != 'Windows':
        print("Warning: Automatic printing is only supported on Windows")
        return False
    
    try:
        # Use os.startfile with 'print' verb to print the document
        # This opens the document with the default print action
        os.startfile(filepath, 'print')
        print(f"Document sent to default printer: {filepath}")
        return True
    except Exception as e:
        print(f"Error printing document: {str(e)}")
        return False


def parse_date_argument(date_arg):
    """Parse date argument from command line.
    
    Supported formats:
    - DD.MM.YY or DD.MM.YYYY - specific date
    - DD.MM - day and month of current year
    - integer (0, -1, +1, etc.) - offset from today
    """
    if not date_arg:
        return datetime.now().date()
    
    # Try to parse as integer offset
    try:
        offset = int(date_arg)
        return datetime.now().date() + timedelta(days=offset)
    except ValueError:
        pass
    
    # Try to parse as date string
    today = datetime.now()
    
    # Try different date formats
    for fmt in ['%d.%m.%Y', '%d.%m.%y', '%d.%m']:
        try:
            if fmt == '%d.%m':
                # For DD.MM format, use current year
                parsed = datetime.strptime(date_arg, fmt)
                return parsed.replace(year=today.year).date()
            else:
                return datetime.strptime(date_arg, fmt).date()
        except ValueError:
            continue
    
    # If nothing worked, raise error
    raise ValueError(
        f"Invalid date format: '{date_arg}'. "
        "Use DD.MM.YY, DD.MM.YYYY, DD.MM, or integer offset (0, -1, +1, etc.)"
    )


def load_meeting_room_emails(filename='meeting_room_emails.txt'):
    """Load meeting room email addresses from file."""
    room_emails = set()
    
    try:
        if os.path.exists(filename):
            with open(filename, 'r', encoding='utf-8') as f:
                for line in f:
                    email = line.strip().lower()
                    if email:  # Skip empty lines
                        room_emails.add(email)
            print(f"Loaded {len(room_emails)} meeting room email(s) from {filename}")
        else:
            print(f"Meeting room emails file '{filename}' not found, no rooms will be filtered")
    except Exception as e:
        print(f"Warning: Failed to load meeting room emails: {str(e)}")
    
    return room_emails


def get_partstat_indicator(partstat):
    """Get Unicode indicator for participant status.
    
    Args:
        partstat: PARTSTAT parameter value from iCalendar (ACCEPTED, DECLINED, NEEDS-ACTION, TENTATIVE, etc.)
    
    Returns:
        Unicode character representing the status
    """
    if not partstat:
        return '○'  # Circle for no response/needs action
    
    partstat_upper = partstat.upper()
    
    if partstat_upper == 'ACCEPTED':
        return '✓'  # Checkmark for accepted
    elif partstat_upper == 'DECLINED':
        return '✗'  # Cross for declined
    elif partstat_upper == 'TENTATIVE':
        return '?'  # Question mark for tentative
    elif partstat_upper == 'DELEGATED':
        return '→'  # Arrow for delegated
    else:  # NEEDS-ACTION or unknown
        return '○'  # Circle for needs action


def resolve_attendee_name(attendee_email, email_to_name):
    """Resolve attendee email to full name from addressbook."""
    # Clean email address
    email = attendee_email.lower().strip()
    
    # Remove mailto: prefix if present
    if email.startswith('mailto:'):
        email = email[7:]
    
    # Try to find in addressbook
    if email in email_to_name:
        return email_to_name[email]
    
    # If not found, return the original email
    return attendee_email


def get_events_for_date(calendar, target_date, tz, email_to_name=None, room_emails=None):
    """Fetch events for a specific date."""
    # Initialize room_emails as empty set if not provided
    if room_emails is None:
        room_emails = set()
    
    # Convert target_date to datetime with timezone
    start_of_day = datetime.combine(target_date, datetime.min.time()).replace(tzinfo=tz)
    end_of_day = start_of_day + timedelta(days=1)
    
    # Fetch events
    events = calendar.date_search(
        start=start_of_day,
        end=end_of_day,
        expand=True
    )
    
    # Parse and sort events
    parsed_events = []
    for event in events:
        try:
            vevent = event.vobject_instance.vevent
            
            # Extract start time
            start_value = vevent.dtstart.value if hasattr(vevent, 'dtstart') else None
            end_value = vevent.dtend.value if hasattr(vevent, 'dtend') else None
            
            # Check if it's an all-day event
            is_all_day = False
            if start_value and not isinstance(start_value, datetime):
                is_all_day = True
            
            # Extract event details
            event_data = {
                'start': start_value,
                'end': end_value,
                'is_all_day': is_all_day,
                'summary': str(vevent.summary.value) if hasattr(vevent, 'summary') else 'Без темы',
                'location': str(vevent.location.value) if hasattr(vevent, 'location') else '',
                'attendees': []
            }
            
            # Extract organizer email
            organizer_email = None
            if hasattr(vevent, 'organizer'):
                organizer_str = str(vevent.organizer.value)
                if organizer_str.startswith('mailto:'):
                    organizer_email = organizer_str[7:].lower()
                else:
                    organizer_email = organizer_str.lower()
            
            # Extract attendees (excluding room emails and organizer) with their roles and participation status
            attendee_data = []  # Will store dicts with email, role, and partstat
            if hasattr(vevent, 'attendee_list'):
                for attendee in vevent.attendee_list:
                    # Extract email from attendee
                    attendee_str = str(attendee.value)
                    if attendee_str.startswith('mailto:'):
                        attendee_email = attendee_str[7:]
                    else:
                        attendee_email = attendee_str
                    
                    # Skip meeting room emails and organizer
                    if attendee_email.lower() not in room_emails and attendee_email.lower() != organizer_email:
                        # Extract ROLE and PARTSTAT parameters if available
                        role = None
                        partstat = None
                        if hasattr(attendee, 'params'):
                            role = attendee.params.get('ROLE', [None])[0]
                            partstat = attendee.params.get('PARTSTAT', [None])[0]
                        attendee_data.append({'email': attendee_email, 'role': role, 'partstat': partstat})
            elif hasattr(vevent, 'attendee'):
                attendee = vevent.attendee
                attendee_str = str(attendee.value)
                if attendee_str.startswith('mailto:'):
                    attendee_email = attendee_str[7:]
                else:
                    attendee_email = attendee_str
                
                # Skip meeting room emails and organizer
                if attendee_email.lower() not in room_emails and attendee_email.lower() != organizer_email:
                    # Extract ROLE and PARTSTAT parameters if available
                    role = None
                    partstat = None
                    if hasattr(attendee, 'params'):
                        role = attendee.params.get('ROLE', [None])[0]
                        partstat = attendee.params.get('PARTSTAT', [None])[0]
                    attendee_data.append({'email': attendee_email, 'role': role, 'partstat': partstat})
            
            # Resolve attendee names from addressbook and separate by role
            required_attendees = []
            optional_attendees = []
            
            for att_data in attendee_data:
                email = att_data['email']
                role = att_data['role']
                partstat = att_data['partstat']
                
                # Resolve name if possible
                resolved_name = None
                if email_to_name:
                    resolved_name = resolve_attendee_name(email, email_to_name)
                    # Check if name was resolved (different from email)
                    if resolved_name == email:
                        resolved_name = None
                
                attendee_info = {'name': resolved_name, 'email': email, 'partstat': partstat}
                
                # Classify by role: OPT-PARTICIPANT is optional, everything else is required
                if role == 'OPT-PARTICIPANT':
                    optional_attendees.append(attendee_info)
                else:
                    required_attendees.append(attendee_info)
            
            # Sort each group alphabetically by name or email
            required_attendees.sort(key=lambda x: x['name'] if x['name'] else x['email'])
            optional_attendees.sort(key=lambda x: x['name'] if x['name'] else x['email'])
            
            event_data['required_attendees'] = required_attendees
            event_data['optional_attendees'] = optional_attendees
            # Keep 'attendees' for backward compatibility (all attendees combined)
            event_data['attendees'] = required_attendees + optional_attendees
            
            # Handle datetime objects for start
            if event_data['start']:
                if isinstance(event_data['start'], datetime):
                    # If it's a datetime, convert to local timezone
                    if event_data['start'].tzinfo is None:
                        event_data['start'] = event_data['start'].replace(tzinfo=tz)
                    else:
                        event_data['start'] = event_data['start'].astimezone(tz)
                else:
                    # If it's a date (all-day event), keep as date
                    pass
            
            # Handle datetime objects for end
            if event_data['end']:
                if isinstance(event_data['end'], datetime):
                    # If it's a datetime, convert to local timezone
                    if event_data['end'].tzinfo is None:
                        event_data['end'] = event_data['end'].replace(tzinfo=tz)
                    else:
                        event_data['end'] = event_data['end'].astimezone(tz)
                else:
                    # If it's a date (all-day event), keep as date
                    pass
            
            # Calculate duration if not all-day
            if not is_all_day and event_data['start'] and event_data['end']:
                duration = event_data['end'] - event_data['start']
                event_data['duration'] = duration
            else:
                event_data['duration'] = None
            
            parsed_events.append(event_data)
        
        except Exception as e:
            print(f"Warning: Failed to parse event: {str(e)}")
            continue
    
    # Sort events by start time
    def sort_key(x):
        if x['start']:
            if isinstance(x['start'], datetime):
                return x['start']
            else:
                # For all-day events, put them at the beginning
                return datetime.combine(x['start'], datetime.min.time()).replace(tzinfo=tz)
        return datetime.max.replace(tzinfo=tz)
    
    parsed_events.sort(key=sort_key)
    
    return parsed_events


def format_time_cell(event):
    """Format time cell with start time, end time, and duration."""
    if event['is_all_day']:
        return 'Весь день'
    
    if not event['start']:
        return ''
    
    # Format start time
    time_str = event['start'].strftime('%H:%M')
    
    # Add end time if available
    if event['end']:
        time_str += ' - ' + event['end'].strftime('%H:%M')
    
    # Add duration if available (on a new line)
    if event['duration']:
        total_seconds = int(event['duration'].total_seconds())
        hours = total_seconds // 3600
        minutes = (total_seconds % 3600) // 60
        
        if hours > 0 and minutes > 0:
            duration_str = f"{hours} ч {minutes} мин"
        elif hours > 0:
            duration_str = f"{hours} ч"
        elif minutes > 0:
            duration_str = f"{minutes} мин"
        else:
            duration_str = "0 мин"
        
        time_str += f"\n({duration_str})"
    
    return time_str


def create_word_document_compact(events, output_filename, target_date, document_title='Расписание'):
    """Create Word document with compact 2-column schedule table.
    
    First column (3 cm): Time range and duration
    Second column (remaining space): Event summary, location, and participants
    """
    doc = Document()
    
    # Set page margins to 1.5 cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # Add title
    title = doc.add_heading(document_title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_paragraph.add_run(target_date.strftime('%d.%m.%Y'))
    date_run.font.size = Pt(12)
    
    doc.add_paragraph()  # Empty line
    
    if not events:
        doc.add_paragraph('На сегодня встреч не запланировано.')
    else:
        # Separate all-day events from timed events
        all_day_events = [e for e in events if e['is_all_day']]
        timed_events = [e for e in events if not e['is_all_day']]
        
        # Create table without header
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        table.allow_autofit = False
        
        # Add all-day events in a single row if any exist
        if all_day_events:
            row_cells = table.add_row().cells
            
            # FIRST COLUMN: "Весь день"
            time_paragraph = row_cells[0].paragraphs[0]
            time_paragraph.add_run('Весь день')
            
            # SECOND COLUMN: All events info separated by newline
            info_paragraph = row_cells[1].paragraphs[0]
            
            for event_idx, event in enumerate(all_day_events):
                if event_idx > 0:
                    # Add blank line between events
                    info_paragraph.add_run('\n')
                
                # Event summary (bold)
                summary_run = info_paragraph.add_run(event['summary'])
                summary_run.bold = True
                
                # Location and participants
                location_and_participants = []
                
                if event['location']:
                    location_and_participants.append(event['location'])
                
                # Collect all attendees (required + optional)
                all_attendees = event.get('required_attendees', []) + event.get('optional_attendees', [])
                
                # Format attendees: use name if available, otherwise email
                if all_attendees:
                    attendee_names = []
                    for attendee in all_attendees:
                        if attendee['name']:
                            attendee_names.append(attendee['name'])
                        else:
                            attendee_names.append(attendee['email'])
                    
                    participants_str = "ответственные " + ', '.join(attendee_names)
                    location_and_participants.append(participants_str)
                
                # Add the second line if there's any content
                if location_and_participants:
                    info_paragraph.add_run('\n')
                    details_run = info_paragraph.add_run(', '.join(location_and_participants))
                    details_run.bold = False
        
        # Add timed events to table
        for event in timed_events:
            row_cells = table.add_row().cells
            
            # FIRST COLUMN: Time range and duration
            time_paragraph = row_cells[0].paragraphs[0]
            
            if event['start'] and event['end']:
                # Add time range on first line
                time_range = f"{event['start'].strftime('%H:%M')} - {event['end'].strftime('%H:%M')}"
                time_paragraph.add_run(time_range)
                
                # Add duration on second line in parentheses
                if event['duration']:
                    total_seconds = int(event['duration'].total_seconds())
                    hours = total_seconds // 3600
                    minutes = (total_seconds % 3600) // 60
                    
                    if hours > 0 and minutes > 0:
                        duration_str = f"{hours} ч {minutes} мин"
                    elif hours > 0:
                        duration_str = f"{hours} ч"
                    elif minutes > 0:
                        duration_str = f"{minutes} мин"
                    else:
                        duration_str = "0 мин"
                    
                    time_paragraph.add_run('\n')
                    time_paragraph.add_run(f"({duration_str})")
            
            # SECOND COLUMN: Summary, location, and participants
            info_paragraph = row_cells[1].paragraphs[0]
            
            # Line 1: Event summary (bold)
            summary_run = info_paragraph.add_run(event['summary'])
            summary_run.bold = True
            
            # Line 2: Location and participants
            location_and_participants = []
            
            # Add location if present (use name only if available, otherwise email)
            if event['location']:
                location_and_participants.append(event['location'])
            
            # Collect all attendees (required + optional)
            all_attendees = event.get('required_attendees', []) + event.get('optional_attendees', [])
            
            # Format attendees: use name if available, otherwise email
            if all_attendees:
                attendee_names = []
                for attendee in all_attendees:
                    # If name exists, use it; otherwise use email
                    if attendee['name']:
                        attendee_names.append(attendee['name'])
                    else:
                        attendee_names.append(attendee['email'])
                
                # Add "ответственные" label before the list
                participants_str = "ответственные " + ', '.join(attendee_names)
                location_and_participants.append(participants_str)
            
            # Add the second line if there's any content
            if location_and_participants:
                info_paragraph.add_run('\n')
                details_run = info_paragraph.add_run(', '.join(location_and_participants))
                details_run.bold = False
        
        # Set column widths
        # A4 page width is 21 cm, with 1.5 cm margins on each side = 18 cm available
        # First column: 3 cm, Second column: 15 cm (18 - 3)
        widths = [Cm(3), Cm(15)]
        
        # Set width for each cell in each column
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    
    # Save document
    doc.save(output_filename)
    print(f"Schedule saved to: {output_filename}")


def main():
    """Main function."""
    try:
        # Parse command line arguments
        parser = argparse.ArgumentParser(
            description='Print daily schedule from CalDAV calendar to Word document.',
            formatter_class=argparse.RawDescriptionHelpFormatter,
            epilog="""
Date format examples:
  -d 0              Today
  -d -1             Yesterday
  -d +1             Tomorrow
  -d 15.11.2025     Specific date (DD.MM.YYYY)
  -d 15.11.25       Specific date (DD.MM.YY)
  -d 15.11          15th November of current year (DD.MM)
"""
        )
        parser.add_argument(
            '-d', '--date',
            type=str,
            default=None,
            help='Date for schedule (default: today). Format: DD.MM.YY, DD.MM.YYYY, DD.MM, or integer offset (0=today, -1=yesterday, +1=tomorrow)'
        )
        parser.add_argument(
            '-p', '--print',
            action='store_true',
            help='Print document to default printer after saving (Windows only)'
        )
        args = parser.parse_args()
        
        # If date is not provided, prompt for input with default value
        date_arg = args.date
        if date_arg is None:
            # Get default date (today) in DD.MM.YY format
            default_date = datetime.now().strftime('%d.%m.%y')
            try:
                user_input = input(f"Введите дату для печати расписания [нажмите Enter для {default_date}]: ").strip()
                # If user just pressed Enter (empty input), use default
                date_arg = user_input if user_input else default_date
            except (EOFError, KeyboardInterrupt):
                # If input is interrupted, use default
                print(f"\nИспользуется дата по умолчанию: {default_date}")
                date_arg = default_date
        
        # Parse target date
        target_date = parse_date_argument(date_arg)
        print(f"Generating schedule for: {target_date.strftime('%d.%m.%Y')}")
        
        # Load configuration
        print("Loading configuration...")
        config = load_config()
        
        # Connect to calendar
        print("Connecting to CalDAV server...")
        calendar = connect_to_calendar(
            config['caldav_url'],
            config['caldav_username'],
            config['caldav_password']
        )
        print(f"Connected to calendar: {calendar.name}")
        
        # Load meeting room emails
        print("Loading meeting room emails...")
        room_emails = load_meeting_room_emails()
        
        # Load contacts from CardDAV (optional)
        email_to_name = {}
        if config['carddav_url']:
            print("Connecting to CardDAV server and loading contacts...")
            email_to_name = load_contacts_from_carddav(
                config['carddav_url'],
                config['carddav_username'],
                config['carddav_password']
            )
        else:
            print("CardDAV URL not configured, skipping contact resolution")
        
        # Get events for target date
        print(f"Fetching events for {target_date.strftime('%d.%m.%Y')}...")
        events = get_events_for_date(calendar, target_date, config['timezone'], email_to_name, room_emails)
        print(f"Found {len(events)} event(s)")
        
        # Generate output filename
        filename = f"{config['filename_prefix']}{target_date.strftime('%d.%m.%y')}.docx"
        output_path = os.path.join(config['output_path'], filename)
        
        # Create Word document
        print("Creating Word document...")
        create_word_document_compact(events, output_path, target_date, config['document_title'])
        
        # Print document if requested
        if args.print:
            print("\nSending document to printer...")
            if not print_document(output_path):
                print("Note: Document was saved but printing failed")
        
        print("Done!")
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return 1
    
    return 0


if __name__ == '__main__':
    exit(main())

